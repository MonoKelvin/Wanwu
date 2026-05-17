# -*- coding: utf-8 -*-
"""Generate Wanwu URSD and DD Word documents."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

DOC_DIR = Path(__file__).resolve().parent.parent
TODAY = date.today().isoformat()

# 全文中文字体（含西文 fallback 同一字体族，避免混用 Calibri/Arial）
FONT_CN = "微软雅黑"
def set_run_chinese_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT_CN
    r = run._element.get_or_add_rPr()
    rfonts = r.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_CN)
    rfonts.set(qn("w:hAnsi"), FONT_CN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)
    rfonts.set(qn("w:cs"), FONT_CN)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_chinese_font(paragraph) -> None:
    for run in paragraph.runs:
        set_run_chinese_font(run)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_CN
    style.font.size = Pt(10.5)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_CN)
    rfonts.set(qn("w:hAnsi"), FONT_CN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)
    rfonts.set(qn("w:cs"), FONT_CN)

    for i in range(1, 10):
        name = f"Heading {i}"
        if name not in doc.styles:
            continue
        hs = doc.styles[name]
        hs.font.name = FONT_CN
        hrpr = hs._element.get_or_add_rPr()
        hrfonts = hrpr.get_or_add_rFonts()
        hrfonts.set(qn("w:ascii"), FONT_CN)
        hrfonts.set(qn("w:hAnsi"), FONT_CN)
        hrfonts.set(qn("w:eastAsia"), FONT_CN)
        hrfonts.set(qn("w:cs"), FONT_CN)


def apply_chinese_font_to_document(doc: Document) -> None:
    """遍历全文段落与表格，统一为中文字体。"""
    for paragraph in doc.paragraphs:
        set_paragraph_chinese_font(paragraph)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_chinese_font(paragraph)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p_el = doc.add_paragraph()
    p_el.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_el.add_run(title)
    set_run_chinese_font(run, size_pt=22, bold=True)

    for line in [subtitle, f"版本：v1.0.0", f"日期：{TODAY}", "编制：MonoStudio"]:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(line)
        set_run_chinese_font(r, size_pt=12)
    doc.add_page_break()


def h(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(text, level=level)
    set_paragraph_chinese_font(para)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    set_paragraph_chinese_font(para)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for par in hdr[i].paragraphs:
            for run in par.runs:
                set_run_chinese_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            t.rows[ri + 1].cells[ci].text = text
            for par in t.rows[ri + 1].cells[ci].paragraphs:
                set_paragraph_chinese_font(par)
    doc.add_paragraph()


def fr_block(doc: Document, fr_id: str, title: str, desc: str, priority: str,
             pre: str, scenario: str, accept: str, dd_ref: str) -> None:
    p(doc, f"{fr_id}：{title}")
    p(doc, f"描述：{desc}")
    p(doc, f"优先级：{priority}")
    p(doc, f"前置条件：{pre}")
    p(doc, f"主成功场景：{scenario}")
    p(doc, f"验收标准：{accept}")
    p(doc, f"关联设计：{dd_ref}")
    p(doc, "")


def build_ursd() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc, "Wanwu（万物）", "用户需求文档")

    # Ch1
    h(doc, "1 文档说明")
    p(doc, "1.1 文档目的")
    p(doc, "本文档描述 Wanwu（万物）v1.0.0 的用户需求，面向产品、开发、测试及 AI 辅助实现。"
       "文档采用编号化、可验收的需求描述，便于追溯与自动化拆解任务。")
    p(doc, "1.2 适用范围")
    p(doc, "适用于 PC 桌面平台（Windows / macOS / Linux）首版；不包含移动端与 Web 独立版。")
    p(doc, "1.3 引用文档")
    p(doc, "源需求：doc/Wanwu（万物）软件需求说明.txt")
    p(doc, "1.4 修订记录")
    table(doc, ["版本", "日期", "说明", "作者"],
          [["v1.0.0", TODAY, "初稿，基于软件需求说明编制", "MonoStudio"]])

    # Ch2
    h(doc, "2 产品概述")
    p(doc, "万物（Wanwu）是一款用于收集与浏览「事物」的桌面应用。事物可以是具象实体（如猫、植物、超跑），"
       "也可以是抽象概念（如设计流派、历史事件、电影）。用户可在统一界面中浏览系统预置的全库内容，"
       "订阅 RSS 资讯流，并创建自建分类与自定义字段，以结构化方式沉淀个人兴趣与知识。")
    p(doc, "价值主张：以高级、简约、流畅的浏览体验，将分散的百科知识、图像细节与冷知识汇聚于本地，"
       "支持离线、可扩展、可个性化的个人收藏体系。")

    # Ch3
    h(doc, "3 术语与概念")
    table(doc, ["术语", "定义"],
          [
              ["物品（Item）", "系统中可被浏览的最小内容单元，含名称、简介、媒体、标签等属性"],
              ["全库（Library）", "系统预置默认内容库，用户不可删除分类，可修改与补充条目"],
              ["自建（Custom）", "用户自行创建的分类与物品，分类名不得与全库重复"],
              ["模块导航", "界面最左侧的一级模块入口：全库、RSS、自建、个人、设置"],
              ["细项列表", "界面中部，展示当前模块下的分类、子分类或条目列表"],
              ["展示页", "界面右侧主区域，展示物品详情或模块主页"],
              ["动态控件组", "用户在自建物品时可自定义的输入控件集合（单行、多行、颜色等）"],
          ])

    # Ch4
    h(doc, "4 目标用户与画像")
    personas = [
        ("P1 知识探索者", "22-35岁", "喜欢百科、冷知识、物种/历史", "希望一站式浏览、离线可用", "中等计算机水平"),
        ("P2 视觉收藏者", "20-40岁", "立绘、UI、家装、工业设计等图像类", "需要大图流畅浏览与分类", "较高审美要求"),
        ("P3 RSS 阅读者", "25-45岁", "关注领域资讯聚合", "希望订阅与本地收藏打通", "熟悉 RSS 概念"),
        ("P4 深度整理者", "28-50岁", "自建专题库、自定义字段", "需要灵活表单与标签体系", "愿意花时间录入"),
    ]
    table(doc, ["画像", "年龄", "动机", "期望", "技能"],
          [list(x) for x in personas])

    # Ch5
    h(doc, "5 用户痛点")
    pains = [
        "信息分散在网页、相册、笔记等多处，难以统一检索与回顾",
        "通用笔记软件缺乏针对「事物」的结构化展示（首屏关键信息 + 滚动细节）",
        "RSS 阅读器与本地图鉴/百科割裂，无法在同一产品内关联",
        "自建收藏时字段固定，无法表达领域特有属性（如色号、排量等）",
        "许多图鉴类应用 UI 陈旧，缺乏高级感与流畅动效",
        "担心隐私：不希望所有收藏强制上传云端",
    ]
    for i, pain in enumerate(pains, 1):
        p(doc, f"痛点{i}：{pain}")

    # Ch6
    h(doc, "6 用户期望")
    expectations = [
        "本地优先存储，数据可控，v1.0 支持加密",
        "界面扁平、现代、黑白灰为主，浅色主题，预留深色变量",
        "全库内容丰富且可纠错补充，但不允许删除系统分类",
        "自建分类智能去重，引导使用全库已有分类",
        "物品详情页参考专业电商/杂志排版，信息层次清晰",
        "适当动效与过渡，不喧宾夺主",
    ]
    for i, e in enumerate(expectations, 1):
        p(doc, f"期望{i}：{e}")

    # Ch7 FR
    h(doc, "7 功能需求")
    p(doc, "本章按模块列出功能需求，编号规则：FR-{模块缩写}-{序号}。")

    h(doc, "7.1 全库模块", 2)
    library_frs = [
        ("FR-LIB-001", "默认分类展示", "系统启动后全库模块展示14个默认一级分类",
         "P0", "已安装应用", "用户进入全库→看到猫、狗、超跑等分类",
         "列表包含：猫、狗、超跑、立绘、UI设计、植物、摩托、家装设计、工业设计、变形金刚、超级英雄动漫、日漫、历史、电影",
         "DD-6.1"),
        ("FR-LIB-002", "子分类浏览", "每个一级分类可展开多种子分类于细项列表",
         "P0", "已选一级分类", "用户点击猫→细项列表显示品种/用途等子类",
         "至少3个子分类可点击；选中后右侧展示对应物品列表或详情", "DD-6.1"),
        ("FR-LIB-003", "物品列表", "细项选中后在展示页或列表区显示物品",
         "P0", "已选子分类", "用户选择子类→看到物品卡片/列表",
         "每项显示缩略图、名称、一行简介；点击进入详情", "DD-7"),
        ("FR-LIB-004", "禁止删除分类", "用户不能删除全库一级/二级分类",
         "P0", "全库模块", "用户尝试删除分类",
         "界面无删除分类入口；API 层拒绝 delete category", "DD-6.1"),
        ("FR-LIB-005", "允许编辑物品", "用户可修改全库物品的名称、简介、详述、标签、图片等",
         "P0", "有编辑权限", "用户打开物品→编辑→保存",
         "修改持久化；刷新后仍生效；记录更新时间", "DD-7"),
        ("FR-LIB-006", "允许补充物品", "用户可在子分类下新增物品",
         "P1", "已选子分类", "用户点击新增→填写表单→保存",
         "新物品出现在列表；字段校验与全库一致", "DD-7"),
        ("FR-LIB-007", "搜索", "支持在全库内按名称、标签搜索",
         "P1", "全库有数据", "用户输入关键词",
         "结果实时过滤；空结果有提示", "DD-6.1"),
        ("FR-LIB-008", "初始数据", "首次安装含大量预置物品",
         "P0", "首次启动", "用户浏览各分类",
         "每类至少有若干条示例物品；信息基本准确", "DD-11"),
        ("FR-LIB-009", "分类图标", "一级分类配有识别性图标",
         "P2", "全库模块", "用户浏览分类导航",
         "每类有 monochrome 图标；风格统一", "DD-12"),
    ]
    for args in library_frs:
        fr_block(doc, *args)

    h(doc, "7.2 RSS 模块", 2)
    rss_frs = [
        ("FR-RSS-001", "默认订阅源", "系统提供若干免费默认 RSS/API 源",
         "P1", "进入 RSS 模块", "用户查看默认源列表",
         "至少3个可启用源；可预览条目", "DD-9"),
        ("FR-RSS-002", "自定义订阅", "用户可通过通用表单添加订阅",
         "P0", "RSS 模块", "用户新建订阅→填 URL/参数→保存",
         "保存成功；出现在订阅列表", "DD-9"),
        ("FR-RSS-003", "条目解析展示", "拉取并解析 RSS/Atom 为可读列表",
         "P0", "有效订阅 URL", "用户打开订阅",
         "显示标题、摘要、时间；可点击进入详情或外链", "DD-9"),
        ("FR-RSS-004", "刷新", "支持手动刷新与定时刷新",
         "P1", "有订阅", "用户点击刷新",
         "显示最新条目；失败有错误提示", "DD-9"),
        ("FR-RSS-005", "转为物品收藏", "可将 RSS 条目收藏到个人收藏",
         "P2", "有条目", "用户收藏条目",
         "出现在个人收藏；保留来源链接", "DD-6.4"),
        ("FR-RSS-006", "编辑订阅", "可修改订阅名称、URL、刷新间隔",
         "P1", "已有订阅", "用户编辑→保存",
         "配置更新生效", "DD-9"),
        ("FR-RSS-007", "删除订阅", "用户可删除自建订阅",
         "P1", "已有订阅", "用户删除→确认",
         "订阅及缓存条目移除", "DD-9"),
    ]
    for args in rss_frs:
        fr_block(doc, *args)

    h(doc, "7.3 自建模块", 2)
    custom_frs = [
        ("FR-CUS-001", "创建分类", "用户可创建自建一级分类",
         "P0", "自建模块", "用户新建分类→输入名称→保存",
         "分类出现在细项列表", "DD-6.3"),
        ("FR-CUS-002", "分类去重", "自建分类不得与全库重名",
         "P0", "输入与全库同名", "用户输入重复名称",
         "提示重复并推荐跳转全库对应分类；阻止保存", "DD-6.3"),
        ("FR-CUS-003", "标准创建表单", "创建物品时提供图片列表、名称、简介、详细说明、标签",
         "P0", "已有自建分类", "用户新增物品",
         "各字段可填可空（名称必填）；支持多图", "DD-8"),
        ("FR-CUS-004", "动态控件组", "用户可为分类或物品定义额外控件",
         "P0", "自建分类", "用户添加控件类型→配置→录入数据",
         "支持单行、多行、颜色；保存后详情页展示", "DD-8"),
        ("FR-CUS-005", "控件类型扩展预留", "架构预留数字、日期、单选等",
         "P2", "v1.0", "—", "v1.0 至少实现三种基础类型", "DD-8"),
        ("FR-CUS-006", "编辑删除物品", "用户可编辑、删除自建物品",
         "P0", "有自建物品", "用户编辑或删除",
         "操作生效；媒体文件同步清理", "DD-8"),
        ("FR-CUS-007", "子分类", "自建分类可设子分类",
         "P1", "有一级分类", "用户添加子分类",
         "细项列表层级展示", "DD-6.3"),
        ("FR-CUS-008", "导入导出预留", "v2.0 支持导出 JSON",
         "P2", "—", "—", "v1.0 不要求", "DD-18"),
    ]
    for args in custom_frs:
        fr_block(doc, *args)

    h(doc, "7.4 个人模块", 2)
    personal_frs = [
        ("FR-PER-001", "基本信息", "用户可设置昵称、头像、简介",
         "P1", "个人模块", "用户编辑资料→保存",
         "展示页主页显示更新内容", "DD-6.4"),
        ("FR-PER-002", "收藏夹", "收藏全库/自建/RSS 物品",
         "P0", "浏览物品", "用户点击收藏",
         "个人模块可列表查看；可取消收藏", "DD-6.4"),
        ("FR-PER-003", "个人主页", "展示用户简介与精选内容",
         "P1", "已设置资料", "进入个人主页",
         "布局简约；支持精选条目配置", "DD-6.4"),
        ("FR-PER-004", "浏览历史", "记录最近查看物品",
         "P2", "有浏览行为", "用户查看历史",
         "按时间倒序；可清空", "DD-6.4"),
    ]
    for args in personal_frs:
        fr_block(doc, *args)

    h(doc, "7.5 设置模块", 2)
    settings_frs = [
        ("FR-SET-001", "应用设置", "语言、数据路径、RSS 刷新间隔等",
         "P0", "设置模块", "用户修改设置",
         "持久化；重启后保留", "DD-6.5"),
        ("FR-SET-002", "关于应用", "显示版本、应用简介",
         "P0", "—", "打开关于",
         "显示 v1.0.0 与 MonoStudio 信息", "DD-6.5"),
        ("FR-SET-003", "赞助支持", "展示赞助方式链接",
         "P2", "—", "用户打开赞助页",
         "外链或内嵌说明可访问", "DD-6.5"),
        ("FR-SET-004", "协议政策", "用户协议、隐私政策",
         "P0", "—", "用户查看协议",
         "文档可读", "DD-6.5"),
        ("FR-SET-005", "开源地址", "展示开源仓库链接",
         "P1", "—", "用户点击",
         "打开系统浏览器至仓库 URL", "DD-6.5"),
        ("FR-SET-006", "数据加密开关", "本地数据库加密选项",
         "P0", "首次启动", "用户设置主密码或启用加密",
         "DB 文件加密；未授权无法读取", "DD-10"),
    ]
    for args in settings_frs:
        fr_block(doc, *args)

    h(doc, "7.6 全局与布局", 2)
    global_frs = [
        ("FR-GLO-001", "三栏布局", "左模块导航、中细项、右展示",
         "P0", "启动应用", "任意模块",
         "三栏始终可见；宽度可拖拽（P1）", "DD-5"),
        ("FR-GLO-002", "模块切换", "五大模块一键切换",
         "P0", "—", "点击左侧导航",
         "切换有过渡动画；状态保持", "DD-5"),
        ("FR-GLO-003", "统一物品详情", "全库/自建物品共用详情页组件",
         "P0", "点击物品", "进入详情",
         "首屏：大图/标题/关键属性；滚动：详述、图集、标签", "DD-7"),
        ("FR-GLO-004", "浅色主题", "v1.0 仅浅色，CSS 变量定义颜色",
         "P0", "—", "—",
         "主色黑白灰；变量名统一 --ww-*", "DD-12"),
        ("FR-GLO-005", "动效", "页面切换、列表、字体过渡",
         "P1", "—", "模块/页面切换",
         "动画时长 200-400ms；可系统减少动效", "DD-13"),
        ("FR-GLO-006", "窗口适配", "支持常见桌面分辨率与窗口缩放",
         "P1", "—", "调整窗口大小",
         "最小窗口 1024x640；三栏比例自适应", "DD-5"),
        ("FR-GLO-007", "错误提示", "操作失败时友好提示",
         "P0", "—", "保存失败/网络错误",
         "Toast 或内联错误文案；不崩溃", "DD-15"),
    ]
    for args in global_frs:
        fr_block(doc, *args)

    # Ch8
    h(doc, "8 布局与交互需求")
    p(doc, "8.1 布局：最左侧固定宽度模块导航（建议 64-80px 图标+文字）；"
       "中部细项列表（建议 240-320px）；右侧展示页占据剩余空间。")
    p(doc, "8.2 视觉：扁平化；控件默认无阴影、无圆角、无边框（分割线除外）；"
       "主色 #111111、次级 #666666、背景 #FFFFFF / #F5F5F5。")
    p(doc, "8.3 交互：单击选中、双击可选快捷打开；列表支持键盘上下选择；"
       "返回保持模块与细项选中状态。")

    # Ch9
    h(doc, "9 物品展示需求")
    p(doc, "9.1 统一详情页结构：Hero 区（主图+标题+关键指标）→ 元信息条 → "
       "滚动区（详细说明、参数表、图集、相关标签）。")
    p(doc, "9.2 排版：高级感字体与留白；正文行高 1.6-1.8；标题层级不超过3级。")
    p(doc, "9.3 v2.0：支持用户自定义展示模板及系统模板（FR 不在 v1.0 验收范围）。")

    # Ch10
    h(doc, "10 数据与隐私需求")
    p(doc, "10.1 v1.0 数据仅存本地 userData 目录。")
    p(doc, "10.2 数据库按分类分文件存储；图片等大文件存 media 目录，库内仅存相对路径。")
    p(doc, "10.3 本地库文件加密存储；密钥由用户主密码或系统派生（见设计文档）。")
    p(doc, "10.4 不收集用户行为上报；外链由用户主动打开。")

    # Ch11 scenarios
    h(doc, "11 用户场景")
    scenarios = [
        ("US-01", "浏览猫咪品种", "作为知识探索者，我想在全库「猫」下浏览各品种，以便了解不同猫咪特征。",
         "进入全库→猫→选择「英国短毛猫」→查看详情首屏与图集。"),
        ("US-02", "修正全库条目", "作为用户，我想纠正某植物描述错误，以便本地库更准确。",
         "全库→植物→选中物品→编辑详述→保存。"),
        ("US-03", "订阅设计资讯", "作为 RSS 阅读者，我想添加 UX 博客 RSS，以便集中阅读。",
         "RSS→新建→填 URL→保存→刷新→阅读列表。"),
        ("US-04", "使用默认 RSS 源", "作为新用户，我想一键启用默认源，以便快速体验。",
         "RSS→默认源列表→启用→查看条目。"),
        ("US-05", "创建摩托专题", "作为收藏者，我想自建「复古摩托」分类并录入车辆。",
         "自建→新建分类→添加物品→填图片与参数。"),
        ("US-06", "避免重复分类", "作为用户，我输入「猫」时系统应引导我去全库。",
         "自建→新建分类「猫」→提示重复→按钮跳转全库。"),
        ("US-07", "自定义颜色字段", "作为深度整理者，我想为手办记录官方色号。",
         "自建→控件组→添加颜色控件→录入→详情展示色块与色值。"),
        ("US-08", "收藏电影条目", "作为用户，我想收藏某部电影以便在个人页回顾。",
         "全库→电影→条目→收藏→个人→收藏夹可见。"),
        ("US-09", "查看个人主页", "作为用户，我想展示简介与精选收藏。",
         "个人→编辑资料→设置精选→主页展示。"),
        ("US-10", "配置加密", "作为隐私敏感用户，我想启用数据库加密。",
         "设置→安全→设置主密码→重启后需解锁。"),
        ("US-11", "浏览超跑", "作为视觉收藏者，我想大图浏览超跑列表。",
         "全库→超跑→列表大图模式→进入详情视差滚动。"),
        ("US-12", "RSS 失效处理", "作为用户，订阅失效时我需看到明确错误。",
         "RSS→失效源→刷新→显示错误与重试。"),
    ]
    for sid, title, story, flow in scenarios:
        p(doc, f"{sid}：{title}")
        p(doc, f"User Story：{story}")
        p(doc, f"场景流：{flow}")
        p(doc, "")

    # Ch12
    h(doc, "12 用户行为模型")
    p(doc, "12.1 典型路径：启动 → 选择模块 → 细项列表浏览 → 打开物品详情 → "
       "（可选）收藏/编辑 → 返回列表。")
    p(doc, "12.2 异常路径：分类重名拦截；RSS 拉取失败；加密未解锁时只读提示；"
       "媒体文件缺失时显示占位图。")
    p(doc, "12.3 频率：浏览>收藏>编辑>自建录入；RSS 刷新为日级。")

    # Ch13 NFR
    h(doc, "13 非功能需求（用户视角）")
    table(doc, ["编号", "类别", "描述", "指标"],
          [
              ["NFR-01", "性能", "列表滚动流畅", "60fps 目标，千级列表虚拟滚动"],
              ["NFR-02", "性能", "冷启动", "<3s 至可交互（SSD）"],
              ["NFR-03", "可用性", "离线可用", "除 RSS 外全功能离线"],
              ["NFR-04", "可用性", "空状态", "每列表有空状态插图与文案"],
              ["NFR-05", "安全", "本地加密", "启用后无明文 DB"],
              ["NFR-06", "可访问性", "键盘导航", "主导航可 Tab 切换"],
          ])

    # Ch14
    h(doc, "14 版本范围与排除项")
    p(doc, "v1.0.0 包含：PC 桌面、五大模块、统一物品页、本地 SQLite、加密、"
       "14 类全库、RSS 基础、自建动态表单、浅色主题。")
    p(doc, "v1.0.0 不包含：云端同步、深色主题切换、用户自定义展示模板、移动端、"
       "自动在线更新（可预留）。")

    # Ch15 appendix
    h(doc, "15 附录")
    h(doc, "15.1 默认全库分类及示例子分类", 2)
    cats = [
        ("猫", "品种, 习性, 护理"),
        ("狗", "品种, 训练, 健康"),
        ("超跑", "品牌, 车型, 赛事"),
        ("立绘", "风格, 画师, 作品"),
        ("UI设计", "移动端, 网页, 设计系统"),
        ("植物", "观叶, 花卉, 多肉"),
        ("摩托", "街车, 复古, 越野"),
        ("家装设计", "北欧, 日式, 工业风"),
        ("工业设计", "家电, 交通, 包装"),
        ("变形金刚", "G1, 电影线, 模型"),
        ("超级英雄动漫", "漫威, DC, 独立（注：原文「超超英」按超级英雄动漫理解）"),
        ("日漫", "少年, 少女, 剧场版"),
        ("历史", "古代, 近代, 人物"),
        ("电影", "华语, 欧美, 动画"),
    ]
    table(doc, ["一级分类", "示例子分类（逗号分隔）"],
          [[a, b] for a, b in cats])

    h(doc, "15.2 需求优先级矩阵", 2)
    table(doc, ["优先级", "含义", "示例 FR"],
          [
              ["P0", "首版必须", "FR-LIB-001, FR-CUS-002, FR-GLO-003"],
              ["P1", "首版应含", "FR-LIB-007, FR-RSS-004, FR-GLO-005"],
              ["P2", "可延期", "FR-PER-004, FR-SET-003"],
          ])

    h(doc, "15.3 原始需求追溯", 2)
    table(doc, ["需求说明行号", "URSD 章节"],
          [
              ["L11-14", "8 布局与交互"],
              ["L16-20", "8、9、13"],
              ["L22-27", "7.1-7.5"],
              ["L29-36", "7、9"],
              ["L39-40", "7.6、13"],
              ["L42-47", "10"],
          ])

    h(doc, "15.4 FR 与 DD 追溯矩阵（摘要）", 2)
    trace = [
        ("FR-LIB-001~008", "DD-6.1, DD-7, DD-10, DD-11"),
        ("FR-RSS-001~007", "DD-9, DD-6.2"),
        ("FR-CUS-001~008", "DD-6.3, DD-8"),
        ("FR-PER-001~004", "DD-6.4, DD-10"),
        ("FR-SET-001~006", "DD-6.5, DD-10"),
        ("FR-GLO-001~005", "DD-5, DD-7, DD-12, DD-13"),
    ]
    table(doc, ["FR 范围", "DD 章节"], trace)

    apply_chinese_font_to_document(doc)
    return doc


def build_dd() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc, "Wanwu（万物）", "详细设计文档")

    h(doc, "1 文档说明")
    p(doc, "本文档为 Wanwu v1.0.0 详细设计，与《用户需求文档 v1.0》配套。"
       "面向开发人员及 AI 编码助手，描述架构、模块、数据、接口与实现要点。")
    table(doc, ["版本", "日期", "说明"],
          [["v1.0.0", TODAY, "初稿"]])

    h(doc, "2 设计目标与原则")
    principles = [
        "单一职责：每个模块、组件、服务只负责一类业务",
        "组件复用：物品卡片、详情页、表单控件跨模块共享",
        "分目录分模块：按 library/rss/custom/personal/settings 划分",
        "可测试：主进程服务与渲染层分离，便于 mock IPC",
        "可扩展：v2.0 模板、深色主题、云同步通过接口预留",
        "免费开源优先：第三方 API、字体、图标优先无版权风险方案",
    ]
    for pr in principles:
        p(doc, f"• {pr}")

    h(doc, "3 系统架构")
    p(doc, "3.1 总体架构")
    p(doc, "采用 Electron 双进程模型：主进程（Node）负责 SQLite、文件 IO、RSS 拉取、加密；"
       "渲染进程（Vue3）负责 UI 与交互，通过 IPC 调用主进程服务。")
    p(doc, "架构分层：Presentation（Vue+PrimeVue+Tailwind）→ Application（Pinia stores）→ "
       "IPC Client → Main Services（db, media, rss, crypto）→ Data（SQLite + 文件系统）。")
    p(doc, "3.2 架构示意（文字版）")
    p(doc, "Renderer: ModuleNav → SubItemList → ItemDetailView；"
       "Main: IPC_Bridge → SQLite_per_category + MediaFileStore + RSS_Fetcher。")

    h(doc, "4 技术栈")
    table(doc, ["层次", "技术", "职责"],
          [
              ["桌面壳", "Electron 28+", "窗口、IPC、系统 API"],
              ["UI 框架", "Vue 3 + TypeScript", "组件化界面"],
              ["组件库", "PrimeVue 4", "表格、对话框等基础控件"],
              ["样式", "Tailwind CSS 3", "工具类 + 设计 Token"],
              ["状态", "Pinia", "模块状态、选中项、用户设置"],
              ["路由", "Vue Router", "模块内子路由"],
              ["数据库", "better-sqlite3", "主进程同步 SQLite"],
              ["RSS", "rss-parser", "主进程解析 Feed"],
              ["构建", "electron-vite", "主/渲染打包"],
          ])

    h(doc, "5 应用布局与路由设计")
    p(doc, "5.1 Shell 布局组件 AppShell：left=ModuleSidebar, center=SubItemPanel, right=ContentRouterOutlet。")
    p(doc, "5.2 路由表：/library/:catId?/:subId?, /rss/:feedId?, /custom/:catId?, /personal, /settings。")
    p(doc, "5.3 Pinia stores：useAppStore, useLibraryStore, useRssStore, useCustomStore, useUserStore。")

    h(doc, "6 模块详细设计")
    h(doc, "6.1 全库模块（library）", 2)
    p(doc, "组件：LibraryView, CategoryTree, ItemGrid, ItemDetail（共享）。")
    p(doc, "流程：加载 categories 表 → 渲染左侧导航选中 → 加载 items → 点击 itemId 路由到详情。")
    p(doc, "IPC：library:listCategories, library:listItems, library:updateItem, library:createItem。")

    h(doc, "6.2 RSS 模块（rss）", 2)
    p(doc, "组件：RssView, FeedList, EntryList, EntryDetail。")
    p(doc, "主进程 RssService：fetchFeed(url) → parse → 存 rss_entries；定时器默认 30min。")

    h(doc, "6.3 自建模块（custom）", 2)
    p(doc, "组件：CustomView, CategoryForm, ItemForm, DynamicFieldBuilder, DynamicFieldRenderer。")
    p(doc, "分类创建前调用 library:checkCategoryNameDuplicate。")

    h(doc, "6.4 个人模块（personal）", 2)
    p(doc, "组件：PersonalHome, ProfileEditor, FavoritesList, HistoryList。")
    p(doc, "数据：user.sqlite 中 profiles, favorites, history 表。")

    h(doc, "6.5 设置模块（settings）", 2)
    p(doc, "组件：SettingsView, GeneralSettings, SecuritySettings, AboutPanel, LegalPanel。")

    h(doc, "7 物品模型与统一展示页")
    p(doc, "7.1 Item 实体（TypeScript 接口）")
    p(doc, "interface Item { id, categoryId, subCategoryId, source: 'library'|'custom'|'rss', "
       "name, summary, description, tags: string[], coverMediaId?, mediaIds[], "
       "customFields?: Record<string, unknown>, createdAt, updatedAt }")
    p(doc, "7.2 详情页区块：ItemDetailHero, ItemMetaBar, ItemDescription, ItemGallery, ItemTags。")
    p(doc, "7.3 列表卡片 ItemCard：缩略图 lazy load，标题单行 ellipsis。")

    h(doc, "8 自建动态表单引擎")
    p(doc, "8.1 字段定义表 custom_field_defs：id, categoryId, type, label, order, config(JSON)。")
    p(doc, "8.2 类型枚举 v1.0：text, textarea, color。")
    p(doc, "8.3 值表 custom_field_values：itemId, fieldId, valueText。")
    p(doc, "8.4 DynamicFieldRenderer 根据 type 映射 PrimeVue InputText | Textarea | ColorPicker。")
    p(doc, "8.5 校验：required、maxLength；颜色格式 #RRGGBB。")

    h(doc, "9 RSS 子系统设计")
    p(doc, "9.1 表 rss_feeds(id, title, url, refreshInterval, lastFetchedAt, enabled)。")
    p(doc, "9.2 表 rss_entries(id, feedId, guid, title, summary, link, publishedAt, raw)。")
    p(doc, "9.3 网络：主进程 node-fetch/https；User-Agent 标识 Wanwu/1.0。")
    p(doc, "9.4 失败重试：指数退避 3 次；缓存上次成功结果。")

    h(doc, "10 数据存储设计")
    p(doc, "10.1 目录结构（userData/wanwu/）：db/, media/, cache/, config.json。")
    p(doc, "10.2 库文件：library_{categoryId}.sqlite, custom.sqlite, rss.sqlite, user.sqlite。")
    p(doc, "10.3 核心表：categories, items, item_media, tags, item_tags, "
       "custom_field_defs, custom_field_values, rss_feeds, rss_entries, favorites, history。")
    p(doc, "10.4 媒体路径：media/{source}/{categoryId}/{itemId}/{filename}；DB 存相对路径。")
    p(doc, "10.5 加密方案（v1.0 选定）：使用 @journeyapps/sqlcipher + better-sqlite3 编译扩展；"
       "首次启用加密时迁移明文库；密钥 = PBKDF2(用户主密码, salt)。")
    p(doc, "10.6 备选：应用层 AES-256-GCM 加密整个 .sqlite 文件（实现简单但无法 SQL 查询加密文件）。")

    h(doc, "11 初始数据导入")
    p(doc, "11.1 流水线：seed JSON（assets/seed/）→ ImportService → 各 library DB。")
    p(doc, "11.2 优先级：① API 脚本（build 时）② AI/爬虫整理 JSON ③ 人工最少。")
    p(doc, "11.3 校验：名称非空、图片 URL 可下载或本地占位、分类 ID 合法。")
    p(doc, "11.4 命令：npm run seed:import -- --category=cat --source=eol")

    h(doc, "12 UI/UX 设计规范")
    p(doc, "12.1 CSS 变量（:root）：--ww-bg #FFFFFF; --ww-bg-subtle #F5F5F5; "
       "--ww-text #111111; --ww-text-muted #666666; --ww-border #E5E5E5; --ww-accent #111111。")
    p(doc, "12.2 字体栈：'Inter', 'Source Han Sans SC', 'PingFang SC', system-ui。")
    p(doc, "12.3 字号：标题 20/16/14px；正文 14px；辅助 12px。")
    p(doc, "12.4 PrimeVue 主题覆盖：border-radius: 0; box-shadow: none; 边框 1px solid var(--ww-border)。")

    h(doc, "13 动画与特效")
    p(doc, "13.1 路由过渡：fade + translateY(8px)，300ms ease-out。")
    p(doc, "13.2 列表：stagger 50ms；详情页 Hero 视差滚动可选。")
    p(doc, "13.3 特效边界：高斯模糊仅用于 Hero 背景；渐变用于分隔；避免过度栅格干扰阅读。")

    h(doc, "14 项目目录结构")
    p(doc, """
wanwu/
├── electron/
│   ├── main.ts
│   ├── preload.ts
│   └── services/
│       ├── database.ts
│       ├── libraryService.ts
│       ├── rssService.ts
│       ├── mediaService.ts
│       └── cryptoService.ts
├── src/
│   ├── app/
│   │   ├── App.vue
│   │   ├── router/
│   │   └── styles/tokens.css
│   ├── modules/
│   │   ├── library/
│   │   ├── rss/
│   │   ├── custom/
│   │   ├── personal/
│   │   └── settings/
│   ├── features/item/
│   └── shared/
├── assets/seed/
└── scripts/seed-import.ts
""")

    h(doc, "15 技术难点与方案")
    table(doc, ["难点", "方案"],
          [
              ["大图性能", "缩略图缓存 + 原图懒加载 + 虚拟列表"],
              ["RSS 跨域", "主进程代理请求，渲染进程不直连"],
              ["加密性能", "SQLCipher 页级加密；热数据 WAL"],
              ["分类去重", "主进程维护全库分类名索引缓存"],
              ["动态表单", "JSON Schema 式 defs + 统一 Renderer"],
              ["初始数据量", "分批次导入 + 进度条 + 可中断"],
          ])

    h(doc, "16 第三方 API 资源清单")
    p(doc, "以下 API 优先免费、开源友好；调用均在主进程 SeedService 或构建脚本中执行；"
       "需遵守各站 Terms 与速率限制；失败时回退 assets/seed/*.json。")

    api_rows = [
        ["猫/狗/物种", "Encyclopedia of Life", "https://eol.org/api/docs", "部分免 Key", "CC/BY", "GET /search.json?q={name}", "name→Item.name, description→description, image→media", "1 req/s"],
        ["物种分类", "ITIS", "https://www.itis.gov/ws/description.html", "无", "美国政府公开数据", "GET /jsonservice/SingleRpt/SingleRpt?searchTerm=", "scientificName, commonName", "合理使用"],
        ["植物", "Trefle", "https://trefle.io/api/v1/species/{id}", "Bearer Token 免费注册", "ODbL", "species?filter[common_name]=", "common_name, image_url, family", "120/min 免费档"],
        ["电影", "OMDb", "https://www.omdbapi.com/?apikey=KEY&t=", "apikey 免费申请", "需注明数据来源", "t=Title", "Title, Plot, Poster, Year→tags", "1000/day"],
        ["百科条目", "Wikipedia REST", "https://zh.wikipedia.org/api/rest_v1/page/summary/{title}", "无", "CC BY-SA", "page/summary", "title, extract, thumbnail", "遵守 UA 政策"],
        ["图片", "Wikimedia Commons", "https://commons.wikimedia.org/w/api.php", "无", "各文件许可不同", "action=query&prop=imageinfo", "url→item_media", "检查每张图许可"],
        ["词典/冷知识", "Free Dictionary API", "https://api.dictionaryapi.dev/api/v2/entries/en/{word}", "无", "开源", "entries[0].meanings", "word, definition→summary", "无硬性限制"],
        ["车辆参考", "NHTSA vPIC", "https://vpic.nhtsa.dot.gov/api/", "无", "美国公共领域", "GetModelsForMake/{make}", "Make, Model→超跑/摩托", "合理使用"],
        ["RSS", "各 Feed URL", "用户/默认源", "—", "—", "rss-parser", "title, link, summary→条目", "主进程拉取"],
    ]
    table(doc, ["领域", "API", "文档URL", "认证", "许可", "示例端点", "字段映射", "限流"],
          api_rows)

    p(doc, "16.1 OMDb 调用示例")
    p(doc, "GET https://www.omdbapi.com/?apikey={KEY}&t=Inception&plot=full → 映射 Item(name, description, cover from Poster)。")
    p(doc, "16.2 EOL 调用示例")
    p(doc, "GET https://eol.org/api/search/1.0.json?q=Felis+catus&page=1 → 取 results[0].id → pages API 取描述与图。")
    p(doc, "16.3 Wikipedia 调用示例")
    p(doc, "GET https://zh.wikipedia.org/api/rest_v1/page/summary/猫 → extract 作 summary，thumbnail.source 下载到 media。")
    p(doc, "16.4 默认 RSS 源建议（免费）")
    p(doc, "可根据领域选用公开 RSS：如 Wikimedia RecentChanges、设计类博客官方 Feed；"
       "内置列表存 rss_feeds 表 isDefault=1。")

    h(doc, "17 测试策略")
    p(doc, "17.1 单元测试：DynamicFieldRenderer 校验、分类去重逻辑、Item 序列化。")
    p(doc, "17.2 API 集成测试：mock fetch 返回 fixture JSON，断言 ImportService 写入 DB 字段正确。")
    p(doc, "17.3 IPC 契约测试：preload 暴露 API 与主进程 handler 名称一致。")
    p(doc, "17.4 E2E：v1.0 暂不强制；可选用 Playwright for Electron 抽冒烟。")

    h(doc, "18 v2.0 预留设计")
    p(doc, "18.1 ItemDisplayTemplate 表与模板编辑器；物品 templateId 字段。")
    p(doc, "18.2 主题：tokens 双套 light/dark，设置页一键切换。")
    p(doc, "18.3 SyncAdapter 接口：push/pull，v1.0 不实现。")

    h(doc, "附录 A 数据库 ER 说明")
    p(doc, "categories 1─N items；items 1─N item_media；items N─M tags；"
       "custom_field_defs 1─N custom_field_values；rss_feeds 1─N rss_entries；"
       "users 1─N favorites（多态 item 引用）。")

    h(doc, "附录 B 关键 IPC 列表")
    ipc = [
        ("library:listCategories", "{}", "Category[]"),
        ("library:listItems", "{ categoryId, subCategoryId? }", "Item[]"),
        ("library:updateItem", "{ item }", "Item"),
        ("custom:checkDuplicate", "{ name }", "{ duplicate, suggestModule }"),
        ("rss:fetchFeed", "{ feedId }", "RssEntry[]"),
        ("media:saveImage", "{ buffer, path }", "{ relativePath }"),
        ("db:unlock", "{ password }", "{ ok }"),
    ]
    table(doc, ["Channel", "Args", "Returns"], ipc)

    h(doc, "附录 C FR→DD 追溯矩阵")
    trace_full = [
        ("FR-LIB-001", "DD-6.1", "默认分类配置 assets/seed/categories.json"),
        ("FR-LIB-002", "DD-6.1", "categories.parentId 树形"),
        ("FR-LIB-003", "DD-7", "ItemGrid + ItemDetail"),
        ("FR-LIB-004", "DD-6.1", "无 deleteCategory API"),
        ("FR-LIB-005", "DD-7, DD-10", "library:updateItem"),
        ("FR-LIB-008", "DD-11, DD-16", "SeedService + API"),
        ("FR-RSS-002", "DD-9", "Feed 表单 + rss_feeds"),
        ("FR-CUS-002", "DD-6.3", "checkDuplicate"),
        ("FR-CUS-004", "DD-8", "DynamicFieldBuilder"),
        ("FR-GLO-003", "DD-7", "features/item"),
        ("FR-SET-006", "DD-10", "SQLCipher"),
        ("FR-GLO-004", "DD-12", "CSS variables"),
    ]
    table(doc, ["FR", "DD 章节", "实现要点"], trace_full)

    apply_chinese_font_to_document(doc)
    return doc


def main() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    ursd_path = DOC_DIR / "Wanwu（万物）用户需求文档 v1.0.docx"
    dd_path = DOC_DIR / "Wanwu（万物）详细设计文档 v1.0.docx"
    ursd = build_ursd()
    dd = build_dd()
    apply_chinese_font_to_document(ursd)
    apply_chinese_font_to_document(dd)
    ursd.save(str(ursd_path))
    dd.save(str(dd_path))
    print(f"Generated: {ursd_path}")
    print(f"Generated: {dd_path}")


if __name__ == "__main__":
    main()
